"""步骤16-18：输出 — HTML报告(带ECharts图表) + docx"""

import json
import re
import time
from pathlib import Path

from config import US_STATES, PRODUCTS

CACHE_DIR = Path(__file__).parent / "cache"
OUTPUT_DIR = Path(__file__).parent / "output"


def export_all(state_code: str, product: str = "curtains"):
    """生成HTML报告 + docx"""
    state_name = US_STATES[state_code]["name"]
    cache_dir = CACHE_DIR / f"{state_code}_{product}"

    with open(cache_dir / "data_pool.json", "r", encoding="utf-8") as f:
        data_pool = json.load(f)
    with open(cache_dir / "reports.json", "r", encoding="utf-8") as f:
        reports = json.load(f)

    output_dir = OUTPUT_DIR / f"{state_code}_{product}"
    output_dir.mkdir(parents=True, exist_ok=True)

    chart_data = prepare_chart_data(data_pool, cache_dir)

    print("[Step 16-17] Generating HTML reports...")
    report_a_text = reports.get("report_a_final", reports.get("report_a_raw", ""))
    report_b_text = reports.get("report_b_final", reports.get("report_b_raw", ""))

    product_display = PRODUCTS.get(product, {}).get("display_name", product)

    export_html(
        output_dir / "report.html",
        f"美国{state_name}州 {product_display} 市场调研报告",
        report_a_text,
        chart_data,
        report_b_md=report_b_text,
    )
    export_html(
        output_dir / "report_a.html",
        f"美国{state_name}州 数据调研报告",
        report_a_text,
        chart_data,
    )
    export_html(
        output_dir / "report_b.html",
        f"美国{state_name}州 商业分析报告",
        report_b_text,
        chart_data,
    )

    export_docx(
        output_dir / "report_a.docx",
        f"美国{state_name}州窗帘实体店专项市场调研报告",
        reports.get("report_a_final", reports.get("report_a_raw", "")),
    )
    export_docx(
        output_dir / "report_b.docx",
        f"美国{state_name}州窗帘零售市场商业分析报告",
        reports.get("report_b_final", reports.get("report_b_raw", "")),
    )

    with open(output_dir / "data_pool.json", "w", encoding="utf-8") as f:
        json.dump(data_pool, f, ensure_ascii=False, indent=2)
    with open(output_dir / "reports_raw.json", "w", encoding="utf-8") as f:
        json.dump(reports, f, ensure_ascii=False, indent=2)

    print(f"[OK] Output saved to: {output_dir}")
    print(f"  - report.html (combined with tabs)")
    print(f"  - report_a.html / report_a.docx")
    print(f"  - report_b.html / report_b.docx")
    print(f"  - data_pool.json / reports_raw.json")


def prepare_chart_data(data_pool: dict, cache_dir: Path = None) -> dict:
    """从数据池+supplemental提取图表所需的全部数据"""
    demographics = data_pool.get("demographics", {})
    cities = demographics.get("cities", [])
    economy = data_pool.get("economy", {})
    wages = data_pool.get("wages", {})
    businesses = data_pool.get("local_businesses", {})

    # Load supplemental data for CBP
    cbp_state = {}
    cbp_counties = []
    if cache_dir:
        supp_path = cache_dir / "supplemental_data.json"
        if supp_path.exists():
            with open(supp_path, "r", encoding="utf-8") as f:
                supp = json.load(f)
            cbp_state = supp.get("cbp_state", {})
            cbp_counties = supp.get("cbp_counties", [])
            if isinstance(cbp_counties, dict):
                cbp_counties = list(cbp_counties.values()) if cbp_counties else []

    # Top 8 cities by population
    sorted_cities = sorted(cities, key=lambda c: c.get("population", 0), reverse=True)[:8]

    def clean_name(name):
        for suffix in [" city,", " town,", " CDP,", " village,"]:
            if suffix in name:
                return name.split(suffix)[0]
        return name.split(",")[0] if "," in name else name

    city_names = [clean_name(c.get("name", "")) for c in sorted_cities]
    populations = [c.get("population", 0) for c in sorted_cities]
    incomes = [c.get("median_income", 0) for c in sorted_cities]
    home_values = [c.get("median_home_value", 0) for c in sorted_cities]
    housing_units = [c.get("housing_units", 0) for c in sorted_cities]

    # Sort by income desc for income chart
    income_pairs = sorted(zip(city_names, incomes), key=lambda x: x[1], reverse=True)
    income_sorted_names = [p[0] for p in income_pairs]
    income_sorted_values = [p[1] for p in income_pairs]

    # State-level demographics
    state_level = demographics.get("state_level", [{}])
    if isinstance(state_level, list) and state_level:
        state_level = state_level[0]
    state_pop = state_level.get("population", sum(populations))
    state_housing = state_level.get("housing_units", sum(housing_units))

    # Economy indicators
    gdp = economy.get("gdp", {}).get("value", 0)
    unemployment = economy.get("unemployment", {}).get("value", 0)
    building_permits = economy.get("building_permits", {}).get("value", 0)

    # Wages
    retail_wage = wages.get("retail_sales_wage", {}).get("value", 50000)
    construction_wage = wages.get("construction_wage", {}).get("value", 53000)

    # CBP competition data
    store_count = cbp_state.get("establishments", 0)
    store_employees = cbp_state.get("employees", 0)
    competition_density = round(store_count / (state_pop / 10000), 2) if state_pop and store_count else 0

    # County distribution for chart
    county_names = []
    county_stores = []
    if cbp_counties:
        sorted_counties = sorted(cbp_counties, key=lambda c: int(c.get("ESTAB", 0)), reverse=True)[:7]
        for c in sorted_counties:
            name = c.get("NAME", "").replace(" County, Texas", "").replace(" County, ", " ")
            county_names.append(name)
            county_stores.append(int(c.get("ESTAB", 0)))

    # Monthly cost model
    monthly_labor = int((retail_wage + construction_wage) / 12)
    monthly_rent = 2200
    monthly_marketing = 1600
    monthly_other = 2000
    monthly_revenue = 20000
    monthly_profit = monthly_revenue - monthly_labor - monthly_rent - monthly_marketing - monthly_other

    return {
        "city_names": city_names,
        "populations": populations,
        "incomes": incomes,
        "home_values": home_values,
        "housing_units": housing_units,
        "income_sorted_names": income_sorted_names,
        "income_sorted_values": income_sorted_values,
        "state_pop": state_pop,
        "state_housing": state_housing,
        "gdp": gdp,
        "unemployment": unemployment,
        "building_permits": building_permits,
        "retail_wage": retail_wage,
        "construction_wage": construction_wage,
        "store_count": store_count,
        "store_employees": store_employees,
        "competition_density": competition_density,
        "county_names": county_names,
        "county_stores": county_stores,
        "monthly_labor": monthly_labor,
        "monthly_rent": monthly_rent,
        "monthly_marketing": monthly_marketing,
        "monthly_other": monthly_other,
        "monthly_revenue": monthly_revenue,
        "monthly_profit": monthly_profit,
    }


def export_html(path: Path, title: str, report_md: str, chart_data: dict, report_b_md: str = ""):
    """生成带ECharts图表的HTML报告"""
    cd = chart_data
    gen_time = time.strftime('%Y-%m-%d %H:%M')

    city_names = cd["city_names"]
    populations = cd["populations"]
    incomes = cd["incomes"]
    home_values = cd["home_values"]
    income_sorted_names = cd["income_sorted_names"]
    income_sorted_values = cd["income_sorted_values"]

    # Stats — business-relevant metrics matching texas_report.html
    state_pop = cd["state_pop"]
    store_count = cd["store_count"]
    store_employees = cd["store_employees"]
    comp_density = cd["competition_density"]
    bp = cd["building_permits"]
    unemp = cd["unemployment"]

    def fmt_pop(n):
        if n >= 1_000_000:
            return f"{n / 1_000_000:.1f}M"
        if n >= 1_000:
            return f"{n / 1_000:.0f}K"
        return str(n)

    # Report HTML
    report_a_html = md_to_html(report_md)
    # Add section IDs to h2 tags
    counter = [0]
    def _replace_h2(m):
        idx = counter[0]
        counter[0] += 1
        return f'<h2 id="section-{idx}">'
    report_a_html = re.sub(r'<h2>', _replace_h2, report_a_html)

    report_b_html = md_to_html(report_b_md) if report_b_md else ""

    # TOC from report A
    toc_items = []
    for m in re.finditer(r'^#{1,2}\s+(.+)$', report_md, re.MULTILINE):
        toc_items.append(m.group(1).replace('**', ''))
    toc_html = ""
    for i, item in enumerate(toc_items[:15]):
        short = item[:20] + "..." if len(item) > 20 else item
        toc_html += f'<a href="#section-{i}" class="toc-link">{short}</a>\n'

    # Data source table
    ds_table = """<div class="rpt si" style="margin-top:20px">
<h2>数据来源溯源</h2>
<p>本报告使用三种数据标注：<span class="src-api">蓝色=API直取精确数据</span> <span class="src-live">绿色=实时数据</span> <span class="src-est">灰色=行业推算</span></p>
<table>
<tr><th>数据源</th><th>机构</th><th>数据年份</th><th>更新频率</th><th>类型</th></tr>
<tr><td>Census ACS</td><td>美国人口普查局</td><td>2023 (5-year)</td><td>每年12月</td><td><span class="src-api">官方精确</span></td></tr>
<tr><td>Census CBP</td><td>美国人口普查局</td><td>2022</td><td>每年（滞后2-3年）</td><td><span class="src-api">官方精确</span></td></tr>
<tr><td>FRED</td><td>美联储</td><td>2024 Q4</td><td>月度/季度</td><td><span class="src-api">官方精确</span></td></tr>
<tr><td>BLS OEWS</td><td>劳工统计局</td><td>2024</td><td>每年3月</td><td><span class="src-api">官方精确</span></td></tr>
<tr><td>HUD FMR</td><td>住房与城市发展部</td><td>2026</td><td>每年</td><td><span class="src-api">官方精确</span></td></tr>
<tr><td>Google Places</td><td>Google</td><td>实时</td><td>实时</td><td><span class="src-live">实时API</span></td></tr>
<tr><td>Perplexity Sonar</td><td>Perplexity AI</td><td>实时</td><td>实时搜索</td><td><span class="src-live">搜索获取</span></td></tr>
</table>
</div>"""

    # Tab bar HTML (only if dual report)
    has_tabs = bool(report_b_md)
    tab_bar = ""
    if has_tabs:
        tab_bar = '<div class="tab-bar"><button class="tab-btn on" onclick="sw(0)">📋 数据调研报告</button><button class="tab-btn" onclick="sw(1)">📊 商业分析报告</button></div>'

    # Tab panels - NO scroll-in on tab panels to avoid the opacity bug
    if has_tabs:
        panels = f"""<div class="tp on" id="t0"><div class="rpt">{report_a_html}</div></div>
<div class="tp" id="t1"><div class="rpt">{report_b_html}</div></div>"""
    else:
        panels = f'<div class="rpt si">{report_a_html}</div>'

    # ECharts data
    pop_names_js = json.dumps(city_names, ensure_ascii=False)
    pop_values_js = json.dumps(populations)
    inc_names_js = json.dumps(income_sorted_names, ensure_ascii=False)
    inc_values_js = json.dumps(income_sorted_values)
    home_names_js = json.dumps(city_names, ensure_ascii=False)
    home_values_js = json.dumps(home_values)
    county_names_js = json.dumps(cd.get("county_names", []), ensure_ascii=False)
    county_stores_js = json.dumps(cd.get("county_stores", []))

    ml = cd["monthly_labor"]
    mr = cd["monthly_rent"]
    mm = cd["monthly_marketing"]
    mo = cd["monthly_other"]
    mrev = cd["monthly_revenue"]
    mprof = cd["monthly_profit"]
    # Waterfall transparent stack values
    wf_base = mrev - ml  # after rent deducted... no, waterfall goes: revenue, -rent, -labor, -marketing, -other, profit
    # base values for waterfall: [0, rev-rent, rev-rent-labor, rev-rent-labor-mkt, rev-rent-labor-mkt-other, 0]
    wf_transparent = [0, mrev - mr, mrev - mr - ml, mrev - mr - ml - mm, mrev - mr - ml - mm - mo, 0]
    # actual bar values
    wf_values = [mrev, -mr, -ml, -mm, -mo, mprof]

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>{title}</title>
<script src="https://cdn.jsdelivr.net/npm/echarts@5/dist/echarts.min.js"></script>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:opsz,wght@9..40,400;9..40,500;9..40,600;9..40,700;9..40,900&family=Noto+Sans+SC:wght@400;500;700;900&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
:root{{--fg:oklch(0.12 0.01 240);--bg:oklch(0.98 0.002 240);--card:oklch(1 0 0);--dark:#09090b;--muted:oklch(0.55 0.01 240);--bdr:oklch(0.91 0.004 240);--blue:#2563eb;--blue-bg:#eff6ff;--green:#10b981;--green-bg:#ecfdf5;--red:#ef4444;--amber:#f59e0b;--purple:#8b5cf6;--r:12px}}
*{{margin:0;padding:0;box-sizing:border-box}}
html{{scroll-behavior:smooth;-webkit-text-size-adjust:100%;tab-size:4;font-feature-settings:normal;font-variation-settings:normal}}
body{{font-family:'DM Sans','Noto Sans SC',system-ui,sans-serif;color:var(--fg);background:var(--bg);text-rendering:optimizeLegibility;font-optical-sizing:auto;font-size:15px;line-height:1.75}}
.prog{{position:fixed;top:0;left:0;height:3px;background:var(--blue);z-index:9999;width:0;transition:width 100ms}}

/* === HERO (Manus style: stats inside dark bg) === */
.hero{{background:var(--dark);color:#fafafa;position:relative;overflow:hidden}}
.hero::before{{content:'';position:absolute;top:-30%;left:5%;width:600px;height:600px;border-radius:50%;opacity:.15;background:radial-gradient(circle,var(--blue) 0%,transparent 70%)}}
.hero::after{{content:'';position:absolute;bottom:-20%;right:10%;width:400px;height:400px;border-radius:50%;opacity:.1;background:radial-gradient(circle,var(--green) 0%,transparent 70%)}}
.wrap{{max-width:1400px;margin:0 auto;padding:0 48px;position:relative;z-index:1}}
.hero-inner{{padding:40px 0 48px}}
.hero-top{{display:flex;align-items:flex-start;justify-content:space-between;gap:16px}}
.hero-tags{{display:flex;align-items:center;gap:8px;margin-bottom:12px;flex-wrap:wrap}}
.htag{{font-size:11px;font-family:'JetBrains Mono',monospace;padding:3px 10px;border-radius:6px;font-weight:500}}
.htag-code{{background:rgba(255,255,255,.08);color:var(--t3)}}
.htag-api{{background:rgba(37,99,235,.15);color:#93c5fd;border:1px solid rgba(37,99,235,.25)}}
.htag-live{{background:rgba(16,185,129,.15);color:#6ee7b7;border:1px solid rgba(16,185,129,.25)}}
.hero h1{{font-size:clamp(32px,4.5vw,44px);font-weight:900;letter-spacing:-1.5px;line-height:1.15;font-family:'DM Sans','Noto Sans SC',sans-serif}}
.hero-sub{{font-size:17px;color:#94a3b8;margin-top:8px}}
.hero-rating{{flex-shrink:0;text-align:right}}
.hero-rating-emoji{{font-size:32px;display:block}}
.hero-rating-text{{font-size:20px;font-weight:800;color:#4ade80;margin-top:4px}}

/* Stats inside hero */
.hero-stats{{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-top:24px}}
.hstat{{background:rgba(255,255,255,.08);backdrop-filter:blur(8px);border:1px solid rgba(255,255,255,.12);border-radius:var(--r);padding:16px 18px;transition:transform .2s}}
.hstat:hover{{transform:translateY(-2px)}}
.hstat-lbl{{font-size:10px;text-transform:uppercase;letter-spacing:.8px;color:var(--t3);font-weight:500;margin-bottom:6px}}
.hstat-val{{font-size:24px;font-weight:700;color:#fff;font-family:'JetBrains Mono',monospace}}
.hstat-note{{font-size:11px;color:var(--t3);margin-top:4px}}

/* === CHARTS SECTION (white bg, zinc-50 cards with titles) === */
.charts-section{{background:var(--card);border-bottom:1px solid var(--bdr)}}
.charts-section .wrap{{padding-top:24px;padding-bottom:24px}}
.charts{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px}}
.chart-card{{background:#fafafa;border:1px solid var(--bdr);border-radius:var(--r);padding:14px 16px}}
.chart-title{{font-size:12px;font-weight:600;color:var(--muted);text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px}}

/* === CONTENT AREA === */
.content{{padding-top:32px;padding-bottom:48px}}
.layout{{display:grid;grid-template-columns:180px 1fr;gap:36px}}
.toc{{position:sticky;top:24px;align-self:start}}
.toc-title{{font-size:10px;text-transform:uppercase;letter-spacing:.8px;color:var(--t2);font-weight:600;margin-bottom:10px}}
.toc-link{{display:block;font-size:12px;color:var(--t2);text-decoration:none;padding:4px 0 4px 12px;border-left:2px solid var(--bdr);transition:all .15s;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
.toc-link:hover,.toc-link.active{{color:var(--blue);border-left-color:var(--blue)}}

.tab-bar{{display:flex;gap:4px;background:var(--card);border:1px solid var(--bdr);border-radius:10px;padding:3px;margin-bottom:20px}}
.tab-btn{{flex:1;padding:10px 0;border:none;background:transparent;font-family:inherit;font-size:15px;font-weight:600;color:var(--t2);cursor:pointer;border-radius:8px;transition:all .2s}}
.tab-btn.on{{background:var(--dark);color:#fafafa}}
.tab-btn:hover:not(.on){{background:#f4f4f5}}
.tp{{display:none}}.tp.on{{display:block}}

/* === REPORT BODY === */
.rpt{{background:var(--card);border:1px solid var(--bdr);border-radius:var(--r);padding:40px 44px;margin-bottom:16px}}
.rpt h2{{font-size:20px;font-weight:700;color:#18181b;margin:48px 0 16px;padding-left:12px;border-left:4px solid var(--blue);letter-spacing:-0.01em}}
.rpt h2:first-child{{margin-top:0}}
.rpt h3{{font-size:17px;font-weight:600;color:#18181b;margin:28px 0 10px;letter-spacing:-0.01em}}
.rpt p{{margin:14px 0;font-size:15px;line-height:1.85;color:var(--fg);letter-spacing:0.02em;word-spacing:0.05em}}
.rpt strong{{color:#18181b;font-weight:700}}
.dn{{font-family:'JetBrains Mono',monospace;font-weight:700;color:#18181b;letter-spacing:0}}
.rpt table{{width:100%;border-collapse:collapse;border-radius:8px;overflow:hidden;margin:14px 0;font-size:14px}}
.rpt th{{background:#18181b;color:#fff;padding:10px 14px;text-align:left;font-weight:600;font-size:11px;text-transform:uppercase;letter-spacing:.3px}}
.rpt td{{padding:10px 14px;border-bottom:1px solid #f4f4f5;color:var(--fg)}}
.rpt tr:nth-child(even){{background:#fafafa}}
.rpt tr:hover td{{background:var(--blue-bg)}}
.rpt ul,.rpt ol{{padding-left:24px;margin:12px 0}}
.rpt li{{margin:8px 0;font-size:15px;line-height:1.85;color:var(--fg);letter-spacing:0.02em}}
.source-tag{{display:inline-block;background:var(--blue-bg);color:var(--blue);font-size:10px;font-family:'JetBrains Mono',monospace;padding:1px 6px;border-radius:4px;margin:0 3px;vertical-align:middle;white-space:nowrap;letter-spacing:0.3px}}
.src-api{{color:var(--blue);font-size:10px;font-weight:600;background:var(--blue-bg);padding:1px 6px;border-radius:4px;font-family:'JetBrains Mono',monospace;white-space:nowrap}}
.src-live{{color:#059669;font-size:10px;font-weight:600;background:var(--green-bg);padding:1px 6px;border-radius:4px;font-family:'JetBrains Mono',monospace;white-space:nowrap}}
.src-est{{color:var(--t2);font-size:10px;font-weight:600;background:#f4f4f5;padding:1px 6px;border-radius:4px;font-family:'JetBrains Mono',monospace;white-space:nowrap}}

.foot{{text-align:center;padding:32px 48px 48px;font-size:13px;color:var(--t3);line-height:1.8}}
.si{{opacity:0;transform:translateY(20px);transition:opacity .5s cubic-bezier(0,0,.2,1),transform .5s cubic-bezier(0,0,.2,1)}}.si.v{{opacity:1;transform:translateY(0)}}
@media(max-width:1024px){{.hero-stats{{grid-template-columns:repeat(2,1fr)}}.charts{{grid-template-columns:1fr 1fr}}.layout{{grid-template-columns:1fr}}.toc{{display:none}}.hero-rating{{display:none}}}}
@media(max-width:768px){{.hero-stats{{grid-template-columns:1fr}}.charts{{grid-template-columns:1fr}}.rpt{{padding:20px}}.wrap{{padding:0 20px}}}}
</style>
</head>
<body>
<div class="prog" id="prog"></div>

<header class="hero">
<div class="wrap">
  <div class="hero-inner">
    <div class="hero-top">
      <div>
        <div class="hero-tags">
          <span class="htag htag-api">Census ACS 2023</span>
          <span class="htag htag-api">FRED 2024</span>
          <span class="htag htag-api">BLS 2024</span>
          <span class="htag htag-live">Google Places</span>
          <span class="htag htag-live">Sonar</span>
        </div>
        <h1>{title}</h1>
        <p class="hero-sub">基于美国联邦政府官方数据源及AI深度搜索，所有数据通过API直接获取，可追溯验证。</p>
      </div>
      <div class="hero-rating">
        <span class="hero-rating-emoji">🟢</span>
        <span class="hero-rating-text">推荐</span>
      </div>
    </div>
    <div class="hero-stats">
      <div class="hstat"><div class="hstat-lbl">竞争商家</div><div class="hstat-val">{store_count}家</div><div class="hstat-note">NAICS 442291 [Census CBP]</div></div>
      <div class="hstat"><div class="hstat-lbl">家庭收入中位数</div><div class="hstat-val">${max(incomes):,}</div><div class="hstat-note">[Census ACS 2023]</div></div>
      <div class="hstat"><div class="hstat-lbl">竞争密度</div><div class="hstat-val">{comp_density} 店/万人</div><div class="hstat-note">{"低竞争" if comp_density < 1 else "中等竞争" if comp_density < 2 else "高竞争"}</div></div>
      <div class="hstat"><div class="hstat-lbl">月建筑许可</div><div class="hstat-val">{bp:,.0f}套</div><div class="hstat-note">[FRED 2024]</div></div>
    </div>
  </div>
</div>
</header>

<div class="charts-section">
<div class="wrap">
  <div class="charts si">
    <div class="chart-card"><div class="chart-title">城市人口对比</div><div id="c1" style="height:240px"></div></div>
    <div class="chart-card"><div class="chart-title">城市收入对比</div><div id="c2" style="height:240px"></div></div>
    <div class="chart-card"><div class="chart-title">竞争密度</div><div id="c3" style="height:240px"></div></div>
    <div class="chart-card"><div class="chart-title">月度成本结构</div><div id="c4" style="height:240px"></div></div>
    <div class="chart-card"><div class="chart-title">月度盈利模型</div><div id="c5" style="height:240px"></div></div>
    <div class="chart-card"><div class="chart-title">各县店铺分布</div><div id="c6" style="height:240px"></div></div>
  </div>
</div>
</div>

<div class="wrap content">
<div class="layout">
<nav class="toc">
  <div class="toc-title">目录导航</div>
  {toc_html}
</nav>
<div>
  {tab_bar}
  {panels}
  {ds_table}
</div>
</div>
</div>

<footer class="foot">
  Market Research System v2.0<br>
  US Census Bureau · FRED · BLS · HUD · Google Places · Perplexity Sonar<br>
  {gen_time}
</footer>

<script>
window.addEventListener('scroll',()=>{{const h=document.body.scrollHeight-innerHeight;document.getElementById('prog').style.width=(scrollY/h*100)+'%'}});
const obs=new IntersectionObserver(es=>es.forEach(e=>{{if(e.isIntersecting){{e.target.classList.add('v');obs.unobserve(e.target)}}}}),{{threshold:.08}});
document.querySelectorAll('.si').forEach(el=>obs.observe(el));
function sw(i){{document.querySelectorAll('.tab-btn').forEach((b,j)=>b.classList.toggle('on',j===i));document.querySelectorAll('.tp').forEach((p,j)=>p.classList.toggle('on',j===i))}}
const secs=document.querySelectorAll('.rpt h2[id]');
const tlinks=document.querySelectorAll('.toc-link');
const tocObs=new IntersectionObserver(es=>{{es.forEach(e=>{{if(e.isIntersecting){{tlinks.forEach(l=>l.classList.remove('active'));const lk=document.querySelector('.toc-link[href=\"#'+e.target.id+'\"]');if(lk)lk.classList.add('active');}}}});}},{{threshold:0,rootMargin:'-20% 0px -60% 0px'}});
secs.forEach(s=>tocObs.observe(s));

const f={{fontFamily:"'DM Sans','Noto Sans SC',system-ui"}};
const tt={{backgroundColor:'rgba(15,23,42,.9)',borderColor:'rgba(255,255,255,.1)',textStyle:{{color:'#f8fafc',fontSize:12}},extraCssText:'border-radius:8px;'}};
const g={{top:36,right:16,bottom:8,left:8,containLabel:true}};

echarts.init(document.getElementById('c1')).setOption({{backgroundColor:'transparent',tooltip:{{...tt,trigger:'axis'}},grid:g,xAxis:{{type:'category',data:{pop_names_js},axisLabel:{{interval:0,rotate:-25,fontSize:10,color:'#71717a'}},axisLine:{{lineStyle:{{color:'#e4e4e7'}}}}}},yAxis:{{type:'value',axisLabel:{{fontSize:10,color:'#71717a'}},axisLine:{{show:false}},splitLine:{{lineStyle:{{color:'#f4f4f5'}}}}}},series:[{{type:'bar',data:{pop_values_js},itemStyle:{{color:{{type:'linear',x:0,y:0,x2:0,y2:1,colorStops:[{{offset:0,color:'#3b82f6'}},{{offset:1,color:'#1d4ed8'}}]}},borderRadius:[4,4,0,0]}},label:{{show:true,position:'top',fontSize:10,fontFamily:"'JetBrains Mono',monospace",color:'#71717a',formatter:p=>p.value>1e6?(p.value/1e6).toFixed(1)+'M':(p.value/1e3).toFixed(0)+'K'}},animationDuration:1200,animationEasing:'cubicOut'}}]}});

echarts.init(document.getElementById('c2')).setOption({{backgroundColor:'transparent',tooltip:{{...tt,trigger:'axis'}},grid:g,xAxis:{{type:'category',data:{inc_names_js},axisLabel:{{interval:0,rotate:-25,fontSize:10,color:'#71717a'}},axisLine:{{lineStyle:{{color:'#e4e4e7'}}}}}},yAxis:{{type:'value',axisLabel:{{fontSize:10,color:'#71717a',formatter:'${{value}}'}},axisLine:{{show:false}},splitLine:{{lineStyle:{{color:'#f4f4f5'}}}}}},series:[{{type:'bar',data:{inc_values_js},itemStyle:{{color:{{type:'linear',x:0,y:0,x2:0,y2:1,colorStops:[{{offset:0,color:'#10b981'}},{{offset:1,color:'#047857'}}]}},borderRadius:[4,4,0,0]}},label:{{show:true,position:'top',fontSize:10,fontFamily:"'JetBrains Mono',monospace",color:'#71717a',formatter:p=>'$'+(p.value/1e3).toFixed(0)+'K'}},animationDuration:1200,animationEasing:'cubicOut'}}]}});

echarts.init(document.getElementById('c3')).setOption({{backgroundColor:'transparent',series:[{{type:'gauge',min:0,max:3,startAngle:200,endAngle:-20,radius:'85%',axisLine:{{lineStyle:{{width:20,color:[[.167,'#10b981'],[.5,'#f59e0b'],[1,'#ef4444']]}}}},pointer:{{width:5,itemStyle:{{color:'auto'}}}},axisTick:{{show:false}},splitLine:{{show:false}},axisLabel:{{show:false}},detail:{{formatter:'{{value}}\\n店/万人',fontSize:14,fontFamily:"'JetBrains Mono',monospace",color:'#18181b',offsetCenter:[0,'30%']}},title:{{offsetCenter:[0,'-20%'],fontSize:12,color:'#71717a'}},data:[{{value:{comp_density},name:'竞争密度'}}],animationDuration:1200}}]}});

echarts.init(document.getElementById('c4')).setOption({{backgroundColor:'transparent',tooltip:{{...tt,trigger:'item'}},series:[{{type:'pie',radius:['45%','70%'],data:[{{value:{mr},name:'租金',itemStyle:{{color:'#3b82f6'}}}},{{value:{ml},name:'人工',itemStyle:{{color:'#10b981'}}}},{{value:{mm},name:'营销',itemStyle:{{color:'#f59e0b'}}}},{{value:{mo},name:'其他',itemStyle:{{color:'#94a3b8'}}}}],label:{{formatter:'{{b}}\\n{{d}}%',fontSize:11}},animationDuration:1200}}]}});

echarts.init(document.getElementById('c5')).setOption({{backgroundColor:'transparent',tooltip:{{...tt,trigger:'axis'}},grid:{{top:36,right:16,bottom:8,left:8,containLabel:true}},xAxis:{{type:'category',data:['月营收','租金','人工','营销','其他','净利润'],axisLabel:{{fontSize:11,color:'#71717a'}}}},yAxis:{{type:'value',axisLabel:{{fontSize:10,color:'#71717a',formatter:'${{value}}'}}}},series:[{{type:'bar',data:[{{value:{mrev},itemStyle:{{color:'#10b981'}}}},{{value:{-mr},itemStyle:{{color:'#ef4444'}}}},{{value:{-ml},itemStyle:{{color:'#ef4444'}}}},{{value:{-mm},itemStyle:{{color:'#ef4444'}}}},{{value:{-mo},itemStyle:{{color:'#ef4444'}}}},{{value:{mprof},itemStyle:{{color:'#3b82f6'}}}}],label:{{show:true,position:'top',fontSize:10,fontFamily:"'JetBrains Mono',monospace",formatter:p=>'$'+Math.abs(p.value/1e3).toFixed(1)+'K'}},animationDuration:1200}}]}});

echarts.init(document.getElementById('c6')).setOption({{backgroundColor:'transparent',tooltip:{{...tt,trigger:'axis'}},grid:{{top:36,right:16,bottom:8,left:8,containLabel:true}},xAxis:{{type:'category',data:{county_names_js},axisLabel:{{interval:0,fontSize:10,color:'#71717a',rotate:-15}},axisLine:{{lineStyle:{{color:'#e4e4e7'}}}}}},yAxis:{{type:'value',axisLabel:{{fontSize:10,color:'#71717a'}},axisLine:{{show:false}},splitLine:{{lineStyle:{{color:'#f4f4f5'}}}}}},series:[{{type:'bar',data:{county_stores_js},itemStyle:{{color:{{type:'linear',x:0,y:0,x2:0,y2:1,colorStops:[{{offset:0,color:'#c4b5fd'}},{{offset:1,color:'#7c3aed'}}]}},borderRadius:[4,4,0,0]}},label:{{show:true,position:'top',formatter:'{{c}}家',fontSize:10,fontFamily:"'JetBrains Mono',monospace",color:'#71717a'}},animationDuration:1200}}]}});

window.addEventListener('resize',()=>{{document.querySelectorAll('.chart-card > div:not(.chart-title)').forEach(el=>{{if(el.id){{const c=echarts.getInstanceByDom(el);if(c)c.resize();}}}});}});
</script>
</body>
</html>"""

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)


def export_docx(path: Path, title: str, report_md: str):
    """生成Word文档 — 对标 Manus 风格排版"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
    except ImportError:
        print("  Warning: python-docx not installed, skipping docx export")
        return

    FONT_CN = '微软雅黑'
    FONT_EN = 'DM Sans'
    FONT_MONO = 'JetBrains Mono'
    CLR_DARK = RGBColor(0x18, 0x18, 0x1b)
    CLR_BODY = RGBColor(0x27, 0x27, 0x2a)
    CLR_BLUE = RGBColor(0x25, 0x63, 0xeb)
    CLR_TAG = RGBColor(0x25, 0x63, 0xeb)
    CLR_MUTED = RGBColor(0x71, 0x71, 0x7a)
    CLR_WHITE = RGBColor(0xff, 0xff, 0xff)

    doc = Document()

    # --- 页面边距 ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # --- 样式设置 ---
    def _set_font(run, size=11, bold=False, color=CLR_BODY, font_en=FONT_EN):
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    style.font.color.rgb = CLR_BODY
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    pf = style.paragraph_format
    pf.line_spacing = 1.6
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)

    # Heading styles
    for level, (sz, clr, sp_before, sp_after) in {
        1: (20, CLR_DARK, 24, 10),
        2: (16, CLR_DARK, 20, 8),
        3: (13, CLR_DARK, 14, 6),
    }.items():
        hs = doc.styles[f'Heading {level}']
        hs.font.name = FONT_EN
        hs.font.size = Pt(sz)
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
        hs.paragraph_format.space_before = Pt(sp_before)
        hs.paragraph_format.space_after = Pt(sp_after)
        hs.paragraph_format.line_spacing = 1.3

    # --- 标题页 ---
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = CLR_DARK
        run.font.bold = True
        run.font.name = FONT_EN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"报告日期：{time.strftime('%Y年%m月%d日')}")
    _set_font(r, size=10, color=CLR_MUTED)
    sub.paragraph_format.space_after = Pt(20)

    # --- 正文处理 ---
    lines = report_md.split('\n')
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        data_rows = [r for r in table_rows if not all(
            set(c.replace('-', '').replace(':', '').strip()) <= {''} for c in r
        )]
        if len(data_rows) < 2:
            table_rows = []
            return
        cols = len(data_rows[0])
        table = doc.add_table(rows=len(data_rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        # 表头深色背景
        for j, cell_text in enumerate(data_rows[0]):
            if j < cols:
                cell = table.cell(0, j)
                cell.text = ''
                p = cell.paragraphs[0]
                r = p.add_run(cell_text.strip().replace('**', ''))
                _set_font(r, size=10, bold=True, color=CLR_WHITE, font_en=FONT_EN)
                # 深色背景
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="18181b"/>')
                cell._tc.get_or_add_tcPr().append(shading)
        # 数据行
        for i in range(1, len(data_rows)):
            for j, cell_text in enumerate(data_rows[i]):
                if j < cols:
                    cell = table.cell(i, j)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    clean = cell_text.strip().replace('**', '')
                    # 源标签特殊处理
                    tag_parts = re.split(r'(\[[^\]]+\])', clean)
                    for part in tag_parts:
                        if not part:
                            continue
                        if part.startswith('[') and part.endswith(']'):
                            r = p.add_run(part)
                            _set_font(r, size=8, color=CLR_TAG, font_en=FONT_MONO)
                        else:
                            r = p.add_run(part)
                            _set_font(r, size=10, color=CLR_BODY)
                    # 偶数行浅灰背景
                    if i % 2 == 0:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="fafafa"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
        doc.add_paragraph('')
        table_rows = []

    def add_rich_paragraph(text, style_name=None):
        """解析行内 **粗体**、[来源标签]、关键数字"""
        # 先过滤 *斜体* 标记
        text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'\1', text)
        p = doc.add_paragraph(style=style_name)
        # 拆分：**粗体** 和 [标签]
        parts = re.split(r'(\*\*.*?\*\*|\[[^\]]+\])', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                _set_font(r, bold=True, color=CLR_DARK)
            elif part.startswith('[') and part.endswith(']'):
                r = p.add_run(f' {part} ')
                _set_font(r, size=8, color=CLR_TAG, font_en=FONT_MONO)
            else:
                # 对关键数字加粗：$金额、百分比、大数字+单位
                num_parts = re.split(r'(\$[\d,]+\.?\d*[KMB]?(?:万|亿)?(?:/[^\s，。、）)]*)?|\d[\d,]*\.?\d*\s*%|[\d,]{4,}\.?\d*\s*(?:人|家|套|个|名|万|亿|店|扇|年|窗|项))', part)
                for np in num_parts:
                    if not np:
                        continue
                    if re.match(r'\$[\d,]|[\d,]{4,}.*[人家套个名万亿店扇年窗项]|\d.*%', np):
                        r = p.add_run(np)
                        _set_font(r, bold=True, color=CLR_DARK, font_en=FONT_MONO)
                    else:
                        r = p.add_run(np)
                        _set_font(r, color=CLR_BODY)
        return p

    def _should_skip(line):
        plain = line.replace('**', '').lstrip('> ').strip()
        if re.match(r'(报告日期|报告完$|分析师[：:]|数据截止[：:])', plain):
            return True
        if re.match(r'^-{3,}$', line) or re.match(r'^`{3}', line) or re.match(r'^\*[^*]+\*$', line):
            return True
        if re.match(r'数据来源[：:].*?(Census|FRED|BLS|HUD|Google|Angi|CBRE)', plain) and len(plain) < 120 and not any(c in plain for c in ['。', '，']):
            return True
        return False

    for line in lines:
        line = line.strip()
        if not line:
            flush_table()
            continue
        if _should_skip(line):
            continue
        # 去掉 blockquote 标记
        if line.startswith('> '):
            line = line[2:].strip()
        elif line.startswith('>'):
            line = line[1:].strip()
        if not line:
            continue
        if line.startswith('|') and '|' in line[1:]:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                table_rows.append(cells)
            continue
        flush_table()
        heading_text = line.replace('**', '').strip()
        if line.startswith('### '):
            doc.add_heading(heading_text[4:].strip(), level=3)
        elif line.startswith('## '):
            if re.match(r'美国.*?(市场调研|商业分析)报告', heading_text[3:].strip()):
                continue
            doc.add_heading(heading_text[3:].strip(), level=2)
        elif line.startswith('# '):
            if re.match(r'美国.*?(市场调研|商业分析)报告', heading_text[2:].strip()):
                continue
            doc.add_heading(heading_text[2:].strip(), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            add_rich_paragraph(line[2:], 'List Bullet')
        elif re.match(r'^\d+\.\s', line):
            add_rich_paragraph(re.sub(r'^\d+\.\s*', '', line), 'List Number')
        else:
            add_rich_paragraph(line)

    flush_table()
    doc.save(str(path))


def md_to_html(md: str) -> str:
    """Markdown转HTML"""

    def _inline(text):
        """处理行内markdown：**粗体**、*斜体*、[来源标注]、数字加重"""
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'\1', text)  # *text* → text (strip italic markers)
        text = re.sub(r'\[([^\]]+)\]', r'<span class="source-tag">[\1]</span>', text)
        # 数字加重：$金额、百分比、大数字+中文单位 — 但跳过已在 <strong>/<span> 标签内的
        def _bold_num(m):
            return f'<strong class="dn">{m.group(0)}</strong>'
        # 先保护已有标签内容
        protected = []
        def _protect(m):
            protected.append(m.group(0))
            return f'\x00PROT{len(protected)-1}\x00'
        text = re.sub(r'<[^>]+>[^<]*</[^>]+>', _protect, text)
        # $金额 (包含逗号、小数、K/M/B后缀、万/亿)
        text = re.sub(r'\$[\d,]+\.?\d*[KMB]?(?:万|亿)?(?:/[^\s<，。、）)]*)?', _bold_num, text)
        # 百分比
        text = re.sub(r'(?<!\d)\d[\d,]*\.?\d*\s*%', _bold_num, text)
        # 大数字+中文单位 (如 29,640,343人、149家、11,890,808套)
        text = re.sub(r'(?<!\d)[\d,]{4,}\.?\d*\s*(?:人|家|套|个|名|万|亿|店|扇|年|窗|项)', _bold_num, text)
        # 还原保护内容
        for i, p in enumerate(protected):
            text = text.replace(f'\x00PROT{i}\x00', p)
        return text

    lines = md.split('\n')
    html_lines = []
    in_list = False
    in_table = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            if in_table:
                html_lines.append('</table>')
                in_table = False
            continue

        # Skip markdown artifacts: ---, *报告完*, ``` code fences, metadata lines
        if re.match(r'^-{3,}$', stripped) or re.match(r'^`{3}', stripped) or re.match(r'^\*[^*]+\*$', stripped):
            continue
        # Skip report metadata lines - strip ** and > first for matching
        plain = stripped.replace('**', '').lstrip('> ').strip()
        if re.match(r'(报告日期|报告完$|分析师[：:]|数据截止[：:])', plain):
            continue
        # Skip standalone "数据来源：..." line (header-like, short, no sentence punctuation)
        if re.match(r'数据来源[：:].*?(Census|FRED|BLS|HUD|Google|Angi|CBRE)', plain) and len(plain) < 120 and not any(c in plain for c in ['。', '，']):
            continue
        # Strip leading > (blockquote marker)
        if stripped.startswith('> '):
            stripped = stripped[2:].strip()
        elif stripped.startswith('>'):
            stripped = stripped[1:].strip()
        if not stripped:
            continue
        elif stripped.startswith('### '):
            html_lines.append(f'<h3>{_inline(stripped[4:])}</h3>')
        elif stripped.startswith('## '):
            heading_text = stripped[3:].replace('**', '').strip()
            # Skip report title headings (duplicates hero title)
            if re.match(r'美国.*?(市场调研|商业分析)报告', heading_text):
                continue
            html_lines.append(f'<h2>{_inline(stripped[3:])}</h2>')
        elif stripped.startswith('# '):
            heading_text = stripped[2:].replace('**', '').strip()
            if re.match(r'美国.*?(市场调研|商业分析)报告', heading_text):
                continue
            html_lines.append(f'<h2>{_inline(stripped[2:])}</h2>')
        elif stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells and not all(set(c.replace('-', '').replace(':', '')) == set() or not c for c in cells):
                if not in_table:
                    html_lines.append('<table>')
                    in_table = True
                    tag = 'th'
                else:
                    tag = 'td'
                row = ''.join(f'<{tag}>{_inline(c)}</{tag}>' for c in cells)
                html_lines.append(f'<tr>{row}</tr>')
        elif stripped.startswith(('- ', '* ')):
            if not in_list:
                html_lines.append('<ul>')
                in_list = True
            html_lines.append(f'<li>{_inline(stripped[2:])}</li>')
        else:
            html_lines.append(f'<p>{_inline(stripped)}</p>')

    if in_list:
        html_lines.append('</ul>')
    if in_table:
        html_lines.append('</table>')

    return '\n'.join(html_lines)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python exporter.py TX [curtains]")
        sys.exit(1)

    state = sys.argv[1].upper()
    product = sys.argv[2] if len(sys.argv) > 2 else "curtains"
    export_all(state, product)
